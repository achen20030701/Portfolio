"""
文档加载器
支持 TXT、PDF、DOCX 格式
"""
from pathlib import Path
from typing import List
from langchain_core.documents import Document


def load_txt(file_path: Path) -> List[Document]:
    """加载 TXT 文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return [Document(
        page_content=content,
        metadata={"source": str(file_path), "type": "txt"}
    )]


def load_pdf(file_path: Path) -> List[Document]:
    """加载 PDF 文件"""
    from PyPDF2 import PdfReader

    reader = PdfReader(str(file_path))
    documents = []
    for i, page in enumerate(reader.pages):
        content = page.extract_text()
        if content:
            documents.append(Document(
                page_content=content,
                metadata={"source": str(file_path), "page": i, "type": "pdf"}
            ))
    return documents


def load_docx(file_path: Path) -> List[Document]:
    """加载 DOCX 文件"""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    content = "\n".join([para.text for para in doc.paragraphs if para.text])
    return [Document(
        page_content=content,
        metadata={"source": str(file_path), "type": "docx"}
    )]


def load_document(file_path: Path) -> List[Document]:
    """
    根据文件类型加载文档

    支持格式：.txt, .pdf, .docx
    """
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return load_txt(file_path)
    elif suffix == ".pdf":
        return load_pdf(file_path)
    elif suffix == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{suffix}")


# 测试代码
if __name__ == "__main__":
    # 测试加载 TXT
    test_file = Path("../../data/test.txt")
    if test_file.exists():
        docs = load_document(test_file)
        print(f"✅ 加载成功，共 {len(docs)} 个文档")
        print(f"  内容预览：{docs[0].page_content[:100]}...")
    else:
        print("❌ 测试文件不存在")
